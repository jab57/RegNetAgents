#!/usr/bin/env python3
"""
Convert markdown manuscript to DOCX format for bioRxiv submission
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown
from pathlib import Path

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to formatted DOCX"""

    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    doc = Document()

    # Set document properties
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Process markdown line by line
    lines = md_content.split('\n')
    in_table = False
    in_code_block = False
    table_data = []

    for line in lines:
        # Skip code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue

        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)

        # Handle tables
        elif line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            # Skip separator lines
            if re.match(r'\|[\s\-:]+\|', line):
                continue
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
        else:
            # End of table
            if in_table and table_data:
                # Create table
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Light Grid Accent 1'

                for i, row_data in enumerate(table_data):
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        cell = row.cells[j]
                        # Clean markdown formatting
                        cell_text = cell_text.replace('**', '')
                        cell.text = cell_text
                        # Bold first row (header)
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

                table_data = []
                in_table = False

            # Handle lists
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                text = clean_markdown_text(text)
                p = doc.add_paragraph(text, style='List Bullet')
            elif re.match(r'^\d+\.', line.strip()):
                text = re.sub(r'^\d+\.\s*', '', line.strip())
                text = clean_markdown_text(text)
                p = doc.add_paragraph(text, style='List Number')

            # Handle regular paragraphs
            elif line.strip():
                text = clean_markdown_text(line.strip())
                if text:
                    p = doc.add_paragraph(text)

            # Handle blank lines
            else:
                if not in_table:
                    doc.add_paragraph()

    # Save document
    doc.save(docx_file)
    print(f"[OK] Converted {md_file} -> {docx_file}")

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    # Bold and italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # Bold italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)      # Bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)          # Italic
    text = re.sub(r'__(.+?)__', r'\1', text)          # Bold
    text = re.sub(r'_(.+?)_', r'\1', text)            # Italic

    # Links
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    # Inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)

    return text

if __name__ == '__main__':
    from docx2pdf import convert as convert_to_pdf

    md_file = Path('biorxiv/preprint_draft.md')
    docx_file = Path('biorxiv/regnetagents_preprint.docx')
    pdf_file = Path('biorxiv/regnetagents_preprint.pdf')

    print("=" * 60)
    print("Converting Markdown to DOCX + PDF for bioRxiv Submission")
    print("=" * 60)

    # Generate DOCX
    print("\n1. Generating DOCX...")
    convert_markdown_to_docx(md_file, docx_file)

    # Generate PDF from DOCX
    print("\n2. Generating PDF...")
    convert_to_pdf(str(docx_file.absolute()), str(pdf_file.absolute()))
    print(f"[OK] Converted {docx_file} -> {pdf_file}")

    print("\n" + "=" * 60)
    print("[SUCCESS] Both files generated!")
    print("=" * 60)
    print(f"\nGenerated files:")
    print(f"  - DOCX: {docx_file.absolute()}")
    print(f"  - PDF:  {pdf_file.absolute()}")
    print("\nReady for bioRxiv submission!")
    print("Upload the PDF file to bioRxiv.")
