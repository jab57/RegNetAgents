#!/usr/bin/env python3
"""
Convert supplementary_material.md to Word document
"""

from docx import Document
from docx.shared import Pt
import re
import os

def convert_to_docx():
    """Convert supplementary material markdown to Word document"""

    print("Converting supplementary_material.md to Word document...")

    # Input and output files
    input_file = "supplementary_material.md"
    output_file = "supplementary_material.docx"

    # Read markdown
    with open(input_file, "r", encoding="utf-8") as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Parse markdown line by line
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines and separators
        if not line or line == '---':
            i += 1
            continue

        # Handle HTML page break directive
        if line.startswith('<div style="page-break-before: always"></div>'):
            doc.add_page_break()
            i += 1
            continue

        # Handle tables
        if line.startswith('|') and line.endswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse and create table
            if len(table_lines) >= 2:  # At least header and separator
                # Parse header
                header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # Skip separator line (the one with ----)
                data_lines = [line for line in table_lines[2:] if line and not line.startswith('|---')]

                # Create table in Word
                table = doc.add_table(rows=1 + len(data_lines), cols=len(header))
                table.style = 'Light Grid Accent 1'

                # Add header
                for col_idx, header_text in enumerate(header):
                    cell = table.rows[0].cells[col_idx]
                    cell.text = header_text
                    # Bold header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Add data rows
                for row_idx, data_line in enumerate(data_lines):
                    cells = [cell.strip() for cell in data_line.split('|')[1:-1]]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < len(header):
                            table.rows[row_idx + 1].cells[col_idx].text = cell_text

            doc.add_paragraph()  # Space after table
            continue

        # Title (# )
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=0)
            i += 1
            continue

        # Heading 1 (## )
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
            i += 1
            continue

        # Heading 2 (### )
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()

        # Simple parsing for bold, italic, and code
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
            elif part:
                p.add_run(part)

        i += 1

    # Save document
    doc.save(output_file)

    print(f"Word document created: {output_file}")
    print(f"Size: {os.path.getsize(output_file) / 1024:.1f} KB")

if __name__ == "__main__":
    convert_to_docx()
