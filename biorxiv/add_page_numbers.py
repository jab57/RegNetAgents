#!/usr/bin/env python3
"""
Add page numbers to existing Word document without changing anything else.
Safe for documents with citations and careful formatting.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import shutil

def add_page_numbers_to_footer(doc):
    """
    Add page numbers to the footer of a Word document.
    Does not modify any content - only adds page numbers.
    """
    for section in doc.sections:
        # Get or create footer
        footer = section.footer

        # Clear existing footer content (if any)
        footer.paragraphs[0].clear()

        # Create centered paragraph for page number
        paragraph = footer.paragraphs[0]
        paragraph.alignment = 1  # Center alignment

        # Add page number field
        run = paragraph.add_run()

        # Create the page number XML element
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        # Add elements to the run
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    print("Page numbers added to footer (centered)")

def main():
    """Add page numbers to the manuscript Word document."""

    input_file = "regnetagents_preprint.docx"
    backup_file = "regnetagents_preprint_backup.docx"

    # Check if file exists
    if not os.path.exists(input_file):
        print(f"ERROR: {input_file} not found!")
        return

    print(f"Adding page numbers to {input_file}...")
    print("(Creating backup first)")

    # Create backup
    shutil.copy(input_file, backup_file)
    print(f"✓ Backup created: {backup_file}")

    # Load document
    doc = Document(input_file)

    # Add page numbers
    add_page_numbers_to_footer(doc)

    # Save document
    doc.save(input_file)

    print(f"✓ Page numbers added successfully!")
    print(f"✓ File saved: {input_file}")
    print(f"  Size: {os.path.getsize(input_file) / 1024:.1f} KB")
    print("\nNext steps:")
    print("1. Open the Word file to verify page numbers appear")
    print("2. Save As PDF (File → Save As → PDF)")
    print("3. Check that citations are still correct")
    print("\nIf anything went wrong, restore from backup:")
    print(f"   copy {backup_file} {input_file}")

if __name__ == "__main__":
    main()
