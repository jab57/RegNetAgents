#!/usr/bin/env python3
"""
Convert markdown manuscript to DOCX format for bioRxiv submission
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from pathlib import Path

def add_page_break_before(paragraph):
    """Add a page break before a paragraph"""
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

def set_keep_with_next(paragraph):
    """Set paragraph to keep with next paragraph (prevents orphaned headings)"""
    pPr = paragraph._element.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)

def set_keep_together(paragraph):
    """Set paragraph to keep lines together (prevents splitting)"""
    pPr = paragraph._element.get_or_add_pPr()
    keepLines = OxmlElement('w:keepLines')
    pPr.append(keepLines)

def add_page_numbers(doc):
    """Add page numbers to the footer of all sections"""
    for section in doc.sections:
        footer = section.footer
        # Clear existing footer content
        footer.paragraphs[0].text = ""

        # Add centered paragraph for page number
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = paragraph.add_run()

        # Create the page number field code
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        # Add field elements to run
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Format the page number
        run.font.size = Pt(10)

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to formatted DOCX"""

    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    doc = Document()

    # Set document properties
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Process markdown line by line
    lines = md_content.split('\n')
    in_table = False
    in_code_block = False
    table_data = []
    last_paragraph = None

    for i, line in enumerate(lines):
        # Handle code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            # Add code line as monospace paragraph with compact formatting
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0  # Single spacing
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(8)  # Smaller font for compactness
            continue

        # Handle HTML page break directive
        if line.strip().startswith('<div style="page-break-before: always"></div>'):
            # Add page break in Word document
            p = doc.add_paragraph()
            add_page_break_before(p)
            continue

        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            set_keep_with_next(p)
            last_paragraph = p
        elif line.startswith('## '):
            # Add page break before major sections (only CONCLUSIONS now)
            heading_text = line[3:].strip()
            if heading_text in ['CONCLUSIONS']:
                p = doc.add_paragraph()  # Empty paragraph for page break
                add_page_break_before(p)
            p = doc.add_heading(line[3:], level=2)
            set_keep_with_next(p)
            last_paragraph = p
        elif line.startswith('### '):
            # Level 3 headings
            p = doc.add_heading(line[4:], level=3)
            set_keep_with_next(p)
            last_paragraph = p
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            set_keep_with_next(p)
            last_paragraph = p

        # Handle tables
        elif line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_data = []
                # If last paragraph is a table caption, keep it with the table
                if last_paragraph and last_paragraph.text.startswith('Table '):
                    set_keep_with_next(last_paragraph)
            # Skip separator lines
            if re.match(r'\|[\s\-:]+\|', line):
                continue
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
        else:
            # End of table
            if in_table and table_data:
                # Create table
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Light Grid Accent 1'

                # Set table to keep together (prevent splitting across pages)
                tblPr = table._element.xpath('./w:tblPr')
                if tblPr:
                    tblPr = tblPr[0]
                else:
                    tblPr = OxmlElement('w:tblPr')
                    table._element.insert(0, tblPr)

                # Add keep together property
                cantSplit = OxmlElement('w:cantSplit')
                tblPr.append(cantSplit)

                for i, row_data in enumerate(table_data):
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        cell = row.cells[j]
                        # Clean markdown formatting
                        cell_text = cell_text.replace('**', '')
                        cell.text = cell_text
                        # Bold first row (header)
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

                table_data = []
                in_table = False
                last_paragraph = None

            # Handle lists
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                text = clean_markdown_text(text)
                p = doc.add_paragraph(text, style='List Bullet')
                last_paragraph = p
            elif re.match(r'^\d+\.', line.strip()):
                # Always keep numbers as text to prevent Word's auto-numbering continuation
                text = line.strip()
                text = clean_markdown_text(text)
                p = doc.add_paragraph(text)
                last_paragraph = p

            # Handle regular paragraphs
            elif line.strip():
                text = clean_markdown_text(line.strip())
                if text:
                    p = doc.add_paragraph(text)
                    last_paragraph = p

            # Handle blank lines
            else:
                if not in_table:
                    doc.add_paragraph()

    # Add page numbers to footer
    add_page_numbers(doc)

    # Save document
    doc.save(docx_file)
    print(f"[OK] Converted {md_file} -> {docx_file}")

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    # First, extract and protect inline code (backticks) and LaTeX math to preserve underscores
    protected_segments = []

    def save_segment(match):
        protected_segments.append(match.group(0))
        return f"XPROTECTEDX{len(protected_segments)-1}XPROTECTEDX"

    # Protect inline code (backticks)
    text = re.sub(r'`([^`]+)`', save_segment, text)

    # Protect LaTeX math blocks (single and double dollar signs)
    text = re.sub(r'\$\$[^\$]+\$\$', save_segment, text)  # Display math $$...$$
    text = re.sub(r'\$[^\$]+\$', save_segment, text)       # Inline math $...$

    # Bold and italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # Bold italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)      # Bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)          # Italic
    text = re.sub(r'__(.+?)__', r'\1', text)          # Bold
    text = re.sub(r'_(.+?)_', r'\1', text)            # Italic

    # Links
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    # Restore protected content (code and LaTeX math - now protected from underscore removal)
    for i, segment in enumerate(protected_segments):
        text = text.replace(f"XPROTECTEDX{i}XPROTECTEDX", segment)

    return text

if __name__ == '__main__':
    from docx2pdf import convert as convert_to_pdf
    import sys

    # Define documents to convert
    documents = [
        {
            'name': 'Main Manuscript',
            'md': Path('preprint_draft.md'),
            'docx': Path('regnetagents_preprint.docx'),
            'pdf': Path('regnetagents_preprint.pdf')
        },
        {
            'name': 'Supplementary Material',
            'md': Path('supplementary_material.md'),
            'docx': Path('supplementary_material.docx'),
            'pdf': Path('supplementary_material.pdf')
        }
    ]

    print("=" * 70)
    print("BioRxiv Manuscript Conversion: Markdown -> DOCX + PDF")
    print("=" * 70)

    all_success = True

    for doc in documents:
        print(f"\n{'='*70}")
        print(f"Processing: {doc['name']}")
        print('='*70)

        # Check if markdown file exists
        if not doc['md'].exists():
            print(f"[SKIP] {doc['md']} not found")
            continue

        try:
            # Generate DOCX
            print(f"\n1. Generating DOCX...")
            convert_markdown_to_docx(doc['md'], doc['docx'])

            # Generate PDF from DOCX
            print(f"\n2. Generating PDF...")
            convert_to_pdf(str(doc['docx'].absolute()), str(doc['pdf'].absolute()))
            print(f"[OK] Converted {doc['docx']} -> {doc['pdf']}")

            print(f"\n[SUCCESS] {doc['name']}")
            print(f"  - DOCX: {doc['docx'].absolute()}")
            print(f"  - PDF:  {doc['pdf'].absolute()}")

        except Exception as e:
            print(f"\n[ERROR] Failed to convert {doc['name']}")
            print(f"  Error: {e}")
            all_success = False

    print("\n" + "=" * 70)
    if all_success:
        print("[SUCCESS] ALL CONVERSIONS COMPLETED!")
    else:
        print("[WARNING] Some conversions failed - check errors above")
    print("=" * 70)
    print("\nReady for bioRxiv submission!")
    print("Upload the PDF files to bioRxiv.")
