#!/usr/bin/env python3
"""
Convert preprint_draft.md to Word document for bioRxiv submission
Uses python-docx which properly handles tables
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def convert_to_docx():
    """Convert markdown manuscript to Word document"""

    print("Converting preprint_draft.md to Word document...")

    # Input and output files
    input_file = "preprint_draft.md"
    output_file = "regnetagents_preprint.docx"

    # Read markdown
    with open(input_file, "r", encoding="utf-8") as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Parse markdown line by line
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines and separators
        if not line or line == '---':
            i += 1
            continue

        # Check if this is a table (starts with |)
        if line.startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse and create table
            if len(table_lines) >= 2:  # At least header and separator
                # Parse header
                header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # Skip separator line (the one with ----)
                data_lines = [line for line in table_lines[2:] if line and not line.startswith('|---')]

                # Create table in Word
                table = doc.add_table(rows=1 + len(data_lines), cols=len(header))
                table.style = 'Light Grid Accent 1'

                # Add header
                for col_idx, header_text in enumerate(header):
                    cell = table.rows[0].cells[col_idx]
                    cell.text = header_text
                    # Bold header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Add data rows
                for row_idx, data_line in enumerate(data_lines):
                    cells = [cell.strip() for cell in data_line.split('|')[1:-1]]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < len(header):
                            # Remove checkmark symbols that break formatting
                            cell_text = cell_text.replace('✓', '[check]')
                            table.rows[row_idx + 1].cells[col_idx].text = cell_text

            doc.add_paragraph()  # Space after table
            continue

        # Title (# )
        if line.startswith('# '):
            p = doc.add_heading(line.replace('# ', ''), level=0)
            i += 1
            continue

        # Heading 1 (## )
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
            i += 1
            continue

        # Heading 2 (### )
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
            i += 1
            continue

        # Regular paragraph
        # Handle bold **text**
        p = doc.add_paragraph()

        # Simple parsing for bold
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part:
                p.add_run(part)

        i += 1

    # Save document
    doc.save(output_file)

    print(f"Word document created: {output_file}")
    print(f"Size: {os.path.getsize(output_file) / 1024:.1f} KB")
    print("\nNow you can:")
    print("1. Open in Word and Save As PDF")
    print("2. Or submit the .docx directly to bioRxiv (they accept Word files!)")

if __name__ == "__main__":
    convert_to_docx()
